"""
Test script to verify Supabase and Gemini API connections
"""
import sys
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

def test_supabase():
    """Test Supabase connection"""
    print("\n" + "="*60)
    print("Testing Supabase Connection...")
    print("="*60)

    try:
        from config import settings
        from supabase import create_client

        print(f"✓ Supabase URL: {settings.SUPABASE_URL}")
        print(f"✓ Supabase Key: {settings.SUPABASE_KEY[:20]}..." if settings.SUPABASE_KEY else "✗ No key found")

        # Create client
        client = create_client(settings.SUPABASE_URL, settings.SUPABASE_KEY)
        print("✓ Supabase client created successfully")

        # Test database connection by querying documents table
        try:
            result = client.table("documents").select("*").limit(1).execute()
            print(f"✓ Database connection successful (documents table exists)")
            print(f"  Found {len(result.data)} records in documents table")
        except Exception as e:
            print(f"⚠ Warning: Could not query documents table: {str(e)}")
            print("  Make sure you've run database_init.sql in Supabase SQL Editor")

        # Test storage connection
        try:
            buckets = client.storage.list_buckets()
            print(f"✓ Storage connection successful")
            print(f"  Found {len(buckets)} storage buckets:")
            for bucket in buckets:
                print(f"    - {bucket['name']}")

            # Check for required buckets
            bucket_names = [b['name'] for b in buckets]
            if 'original-documents' not in bucket_names:
                print("  ⚠ Warning: 'original-documents' bucket not found")
            if 'completed-documents' not in bucket_names:
                print("  ⚠ Warning: 'completed-documents' bucket not found")

        except Exception as e:
            print(f"⚠ Warning: Could not access storage: {str(e)}")

        print("\n✅ Supabase connection test PASSED")
        return True

    except Exception as e:
        print(f"\n❌ Supabase connection test FAILED: {str(e)}")
        return False


def test_gemini():
    """Test Google Gemini API connection"""
    print("\n" + "="*60)
    print("Testing Google Gemini API Connection...")
    print("="*60)

    try:
        from config import settings
        import google.generativeai as genai

        print(f"✓ Gemini API Key: {settings.GEMINI_API_KEY[:20]}..." if settings.GEMINI_API_KEY else "✗ No key found")

        # Configure Gemini
        genai.configure(api_key=settings.GEMINI_API_KEY)
        print("✓ Gemini API configured")

        # Create model
        model = genai.GenerativeModel('gemini-2.5-flash')
        print("✓ Gemini model created (gemini-2.5-flash)")

        # Test generation with simple prompt
        print("\nTesting text generation...")
        response = model.generate_content("Say 'Hello, I am working!' in exactly those words.")
        result_text = response.text.strip()

        print(f"✓ Generated response: {result_text}")

        # Test with a more complex prompt (placeholder extraction simulation)
        print("\nTesting placeholder extraction...")
        test_doc = "This is an agreement for [COMPANY_NAME] and [EMPLOYEE_NAME]."
        test_prompt = f"""Extract placeholders from this text and return as JSON:
{test_doc}

Return format: [{{"name": "Company Name", "placeholder": "[COMPANY_NAME]"}}]"""

        response = model.generate_content(test_prompt)
        print(f"✓ Placeholder extraction test response received")
        print(f"  Response preview: {response.text[:100]}...")

        print("\n✅ Gemini API connection test PASSED")
        return True

    except Exception as e:
        print(f"\n❌ Gemini API connection test FAILED: {str(e)}")
        print("\nTroubleshooting:")
        print("  1. Check your API key is correct in .env")
        print("  2. Visit https://makersuite.google.com/app/apikey to verify/regenerate")
        print("  3. Ensure you have API quota available")
        return False


def test_document_processing():
    """Test document processing capabilities"""
    print("\n" + "="*60)
    print("Testing Document Processing...")
    print("="*60)

    try:
        from docx import Document
        import io

        # Create a simple test document
        doc = Document()
        doc.add_paragraph("Test Document")
        doc.add_paragraph("Employee: [EMPLOYEE_NAME]")
        doc.add_paragraph("Start Date: [START_DATE]")

        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        print("✓ Test document created in memory")

        # Test extraction
        from services.document_service import document_service

        text = document_service.extract_text_from_docx(doc_bytes.getvalue())
        print(f"✓ Text extraction successful")
        print(f"  Extracted text: {text[:100]}...")

        # Test placeholder detection
        from services.gemini_service import gemini_service

        print("\nTesting Gemini placeholder extraction on test document...")
        placeholders = gemini_service.extract_placeholders(text)

        print(f"✓ Found {len(placeholders)} placeholders:")
        for field in placeholders[:3]:  # Show first 3
            print(f"  - {field['name']}: {field['placeholder']} (type: {field['type']})")

        print("\n✅ Document processing test PASSED")
        return True

    except Exception as e:
        print(f"\n❌ Document processing test FAILED: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Run all tests"""
    print("\n" + "="*60)
    print("LegalDoc Filler - Connection Tests")
    print("="*60)

    results = {
        "supabase": test_supabase(),
        "gemini": test_gemini(),
        "document_processing": test_document_processing()
    }

    print("\n" + "="*60)
    print("Test Summary")
    print("="*60)

    all_passed = True
    for test_name, passed in results.items():
        status = "✅ PASSED" if passed else "❌ FAILED"
        print(f"{test_name.upper()}: {status}")
        if not passed:
            all_passed = False

    print("="*60)

    if all_passed:
        print("\n🎉 All tests passed! Your setup is ready.")
        print("\nNext steps:")
        print("  1. Run the backend: python main.py")
        print("  2. Run the frontend: cd ../frontend && npm run dev")
        print("  3. Visit http://localhost:3000")
        return 0
    else:
        print("\n⚠️  Some tests failed. Please check the errors above.")
        print("\nCommon fixes:")
        print("  - Verify .env file has correct credentials")
        print("  - Run database_init.sql in Supabase SQL Editor")
        print("  - Create storage buckets in Supabase Storage")
        print("  - Check Gemini API key is valid")
        return 1


if __name__ == "__main__":
    sys.exit(main())
