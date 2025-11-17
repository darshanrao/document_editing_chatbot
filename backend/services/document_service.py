from docx import Document
from docx.shared import RGBColor
from typing import List, Dict, Optional, Tuple
import io
from datetime import datetime
from utils.database import db
from services.gemini_service import gemini_service
import re
import mammoth
from docx_parser_converter.docx_to_html.docx_to_html_converter import DocxToHtmlConverter


class DocumentService:
    def __init__(self):
        self.bucket_original = "original-documents"
        self.bucket_completed = "completed-documents"

    def extract_text_from_docx(self, file_data: bytes) -> str:
        """Extract text content from a .docx file"""
        try:
            doc = Document(io.BytesIO(file_data))
            full_text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)

            return '\n'.join(full_text)
        except Exception as e:
            raise Exception(f"Failed to extract text from document: {str(e)}")

    async def process_document(self, document_id: str, file_data: bytes) -> Dict[str, any]:
        """
        Process a document: extract text, identify placeholders, create fields.
        This is the main processing pipeline.
        """
        try:
            # Update status to processing
            db.update_document_status(document_id, "processing")

            # Step 1: Extract text content
            text_content = self.extract_text_from_docx(file_data)
            db.update_document_content(document_id, text_content)

            # Step 2: Use Gemini to identify placeholders
            placeholders = gemini_service.extract_placeholders(text_content)

            if not placeholders:
                raise Exception("No placeholders found in the document")

            # Step 3: Create field records in database
            for field_data in placeholders:
                db.create_field(
                    document_id=document_id,
                    name=field_data["name"],
                    placeholder=field_data["placeholder"],
                    field_type=field_data.get("type", "text"),
                    order=field_data["order"]
                )

            # Step 4: Update document status to ready
            db.update_document_status(document_id, "ready")

            return {
                "success": True,
                "total_fields": len(placeholders),
                "status": "ready"
            }

        except Exception as e:
            db.update_document_status(document_id, "error")
            raise Exception(f"Document processing failed: {str(e)}")

    def get_document_preview(self, document_id: str) -> str:
        """
        Generate HTML preview of the document with current field values.
        Uses Mammoth to convert .docx to HTML with proper formatting.
        """
        document = db.get_document(document_id)
        if not document:
            raise Exception("Document not found")

        try:
            # Get original file from storage
            file_path = document.get("file_path", "")
            if file_path:
                file_data = db.download_file(self.bucket_original, file_path)

                # Convert .docx to HTML using Mammoth
                result = mammoth.convert_to_html(io.BytesIO(file_data))
                html_content = result.value

                # Replace placeholders with values
                fields = db.get_fields(document_id)
                for field in fields:
                    placeholder = field["placeholder"]
                    value = field.get("value")

                    if value:
                        # Wrap filled values in a span for styling
                        html_content = html_content.replace(
                            placeholder,
                            f'<span class="filled-field" style="background-color: #d1fae5; color: #047857; padding: 2px 6px; border-radius: 4px; font-weight: 500;">{value}</span>'
                        )
                    else:
                        # Wrap pending placeholders in a span for styling
                        html_content = html_content.replace(
                            placeholder,
                            f'<span class="pending-field" style="background-color: #fef3c7; color: #92400e; padding: 2px 6px; border-radius: 4px; font-weight: 500; border: 1px solid #fbbf24;">{placeholder}</span>'
                        )

                return html_content
            else:
                # Fallback to text content if file not in storage
                content = document.get("original_content", "")
                fields = db.get_fields(document_id)

                for field in fields:
                    placeholder = field["placeholder"]
                    value = field.get("value")
                    if value:
                        content = content.replace(placeholder, value)

                # Convert plain text to HTML
                return f"<pre style='white-space: pre-wrap; font-family: inherit;'>{content}</pre>"

        except Exception as e:
            print(f"Error generating HTML preview: {e}")
            # Fallback to plain text
            content = document.get("original_content", "")
            return f"<pre style='white-space: pre-wrap; font-family: inherit;'>{content}</pre>"

    def get_completed_document_preview(self, document_id: str) -> str:
        """
        Generate HTML preview of the completed document.
        Uses docx-parser-converter to convert the completed .docx to HTML with formatting preservation.
        """
        document = db.get_document(document_id)
        if not document:
            raise Exception("Document not found")

        try:
            # Try to get completed document from storage first
            completed_file_path = f"{document_id}/completed.docx"
            try:
                file_data = db.download_file(self.bucket_completed, completed_file_path)
                print(f"✓ Using completed document from storage for preview")
            except Exception:
                # Completed document not found, generate it on-demand
                print(f"⚠ Completed document not in storage, generating for preview...")
                original_file_data = db.download_file(
                    self.bucket_original,
                    f"{document_id}/original.docx"
                )
                file_data = self.generate_completed_document(document_id, original_file_data)

            # Convert .docx to HTML using docx-parser-converter (preserves formatting/indentation)
            converter = DocxToHtmlConverter(file_data, use_default_values=True)
            html_content = converter.convert_to_html()

            return html_content

        except Exception as e:
            print(f"Error generating completed document preview: {e}")
            raise Exception(f"Failed to generate preview: {str(e)}")

    def generate_completed_document(self, document_id: str, original_file_data: bytes) -> bytes:
        """
        Generate a completed .docx document with all placeholders filled in.
        Returns the completed document as bytes.
        """
        try:
            # Load the original document
            doc = Document(io.BytesIO(original_file_data))
            fields = db.get_fields(document_id)

            # Create a mapping of placeholder -> value
            replacements = {}
            for field in fields:
                if field.get("value"):
                    replacements[field["placeholder"]] = field["value"]

            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        # Replace in the paragraph
                        paragraph.text = paragraph.text.replace(placeholder, value)

                        # Optional: Highlight filled values in yellow
                        for run in paragraph.runs:
                            if any(value in run.text for value in replacements.values()):
                                run.font.highlight_color = None  # Or use RGBColor for background

            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, value in replacements.items():
                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, value)

            # Save to bytes
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output.getvalue()

        except Exception as e:
            raise Exception(f"Failed to generate completed document: {str(e)}")

    def get_context_for_field(self, document_content: str, placeholder: str,
                              context_chars: int = 200) -> str:
        """
        Get surrounding context for a placeholder to help Gemini generate better questions.
        """
        try:
            index = document_content.find(placeholder)
            if index == -1:
                return ""

            # Get text before and after the placeholder
            start = max(0, index - context_chars)
            end = min(len(document_content), index + len(placeholder) + context_chars)

            context = document_content[start:end]
            return context.strip()
        except Exception:
            return ""

    async def upload_original_document(self, document_id: str, file_data: bytes) -> str:
        """Upload original document to Supabase Storage"""
        file_path = f"{document_id}/original.docx"
        db.upload_file(self.bucket_original, file_path, file_data)
        return file_path

    async def upload_completed_document(self, document_id: str, file_data: bytes) -> str:
        """Upload completed document to Supabase Storage"""
        file_path = f"{document_id}/completed.docx"
        db.upload_file(self.bucket_completed, file_path, file_data)
        return file_path

    def get_completion_summary(self, document_id: str) -> Dict[str, any]:
        """Generate completion summary for a document"""
        document = db.get_document(document_id)
        fields = db.get_fields(document_id)

        if not document:
            raise Exception("Document not found")

        total_fields = len(fields)
        completed_fields = sum(1 for f in fields if f["status"] == "filled")

        # Calculate completion time
        created_at = datetime.fromisoformat(document["created_at"].replace('Z', '+00:00'))
        completed_at = document.get("completed_at")

        if completed_at:
            completed_at = datetime.fromisoformat(completed_at.replace('Z', '+00:00'))
            duration = completed_at - created_at
            completion_time = f"{int(duration.total_seconds() / 60)} minutes"
        else:
            completion_time = "In progress"

        return {
            "filename": document["filename"],
            "fieldsCompleted": completed_fields,
            "totalFields": total_fields,
            "completionTime": completion_time
        }


# Singleton instance
document_service = DocumentService()
